"""Report generation — PDF, Excel, CSV, Word exports."""

from __future__ import annotations

import csv
import io
from datetime import UTC, datetime
from typing import Any

from docx import Document as DocxDocument
from openpyxl import Workbook


class ReportEngine:
    SUPPORTED_FORMATS = ("pdf", "excel", "csv", "word")

    def generate(
        self,
        report_type: str,
        data: dict[str, Any],
        fmt: str = "excel",
        title: str | None = None,
    ) -> tuple[bytes, str, str]:
        report_title = title or f"Mahakosh {report_type.replace('_', ' ').title()}"
        if fmt == "csv":
            return self._to_csv(data, report_title), "text/csv", f"{report_type}.csv"
        if fmt == "excel":
            return self._to_excel(data, report_title), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", f"{report_type}.xlsx"
        if fmt == "word":
            return self._to_word(data, report_title), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{report_type}.docx"
        if fmt == "pdf":
            return self._to_pdf(data, report_title), "application/pdf", f"{report_type}.pdf"
        raise ValueError(f"Unsupported format: {fmt}")

    def _flatten_rows(self, data: dict[str, Any], report_type: str) -> list[dict[str, Any]]:
        if report_type == "gst_summary":
            liability = data.get("liability", data.get("collections", {}))
            return [
                {"metric": "Output CGST", "value": liability.get("output_tax", {}).get("cgst", liability.get("output_cgst", 0))},
                {"metric": "Output SGST", "value": liability.get("output_tax", {}).get("sgst", liability.get("output_sgst", 0))},
                {"metric": "Output IGST", "value": liability.get("output_tax", {}).get("igst", liability.get("output_igst", 0))},
                {"metric": "Input Tax Total", "value": liability.get("input_tax", {}).get("total", liability.get("input_tax_total", 0))},
                {"metric": "Net Liability", "value": liability.get("net_liability", 0)},
            ]
        if report_type in ("vendor_ledger", "purchase_register"):
            return data.get("top_vendors", data.get("vendors", []))
        if report_type in ("sales_register", "customer_report"):
            return data.get("top_customers", data.get("customers", []))
        if report_type == "executive_summary":
            return [{"metric": k, "value": v} for k, v in data.items() if isinstance(v, (int, float, str))]
        if report_type == "financial_summary":
            return [{"metric": k, "value": v} for k, v in data.get("summary", data).items()]
        rows = data.get("rows", data.get("items", []))
        if isinstance(rows, list) and rows:
            return rows
        return [{"key": k, "value": str(v)[:500]} for k, v in data.items() if not isinstance(v, (dict, list))]

    def _to_csv(self, data: dict, title: str) -> bytes:
        rows = self._flatten_rows(data, data.get("report_type", "generic"))
        if not rows:
            rows = [{"message": "No data available"}]
        output = io.StringIO()
        writer = csv.DictWriter(output, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)
        return output.getvalue().encode("utf-8-sig")

    def _to_excel(self, data: dict, title: str) -> bytes:
        wb = Workbook()
        ws = wb.active
        ws.title = title[:31]
        ws.append([title])
        ws.append([f"Generated: {datetime.now(UTC).isoformat()}"])
        ws.append([])

        rows = self._flatten_rows(data, data.get("report_type", "generic"))
        if rows:
            headers = list(rows[0].keys())
            ws.append(headers)
            for row in rows:
                ws.append([row.get(h) for h in headers])

        insights = data.get("insights", {})
        if insights:
            ws.append([])
            ws.append(["Insights"])
            for category in ("observations", "recommendations", "warnings"):
                for item in insights.get(category, []):
                    ws.append([category, item.get("text", str(item))])

        buf = io.BytesIO()
        wb.save(buf)
        return buf.getvalue()

    def _to_word(self, data: dict, title: str) -> bytes:
        doc = DocxDocument()
        doc.add_heading(title, 0)
        doc.add_paragraph(f"Generated: {datetime.now(UTC).strftime('%d %b %Y %H:%M UTC')}")

        rows = self._flatten_rows(data, data.get("report_type", "generic"))
        if rows:
            table = doc.add_table(rows=1, cols=len(rows[0]))
            hdr = table.rows[0].cells
            for i, key in enumerate(rows[0]):
                hdr[i].text = str(key)
            for row in rows:
                cells = table.add_row().cells
                for i, key in enumerate(rows[0]):
                    cells[i].text = str(row.get(key, ""))

        insights = data.get("insights", {})
        for category in ("observations", "recommendations", "warnings", "opportunities"):
            items = insights.get(category, [])
            if items:
                doc.add_heading(category.title(), level=2)
                for item in items:
                    doc.add_paragraph(item.get("text", str(item)), style="List Bullet")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def _to_pdf(self, data: dict, title: str) -> bytes:
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        y = 50
        page.insert_text((50, y), title, fontsize=16)
        y += 30
        page.insert_text((50, y), f"Generated: {datetime.now(UTC).strftime('%d %b %Y')}", fontsize=10)
        y += 25

        rows = self._flatten_rows(data, data.get("report_type", "generic"))
        for row in rows[:40]:
            line = " | ".join(f"{k}: {v}" for k, v in row.items())
            page.insert_text((50, y), line[:120], fontsize=9)
            y += 14
            if y > 750:
                page = doc.new_page()
                y = 50

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
